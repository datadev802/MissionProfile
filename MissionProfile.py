# -*- coding: utf-8 -*-

'''
文件名称：MissionProfile.py
作用：任务剖面构建
作者：杨竣铎
创建日期：2022-05-29
修改日期：2022-06-02
''' 

# 引入系统库
from cgi import test
from operator import index
import os
import sys
from tracemalloc import start
from webbrowser import get
from collections import OrderedDict
from h11 import Data
# 引入三方库
import numpy as np
import pandas as pd
from pandas import DataFrame
from simplejson import OrderedDict
import xlrd
from xlrd.book import Book
import graphviz
from graphviz import Digraph
import networkx as nx
from networkx import DiGraph
import docx
from docx.document import Document
from docx.table import Table
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docx.text.parfmt import ParagraphFormat
# 引入自定义库
from StdColor import my_color_dict
# 全局变量定义
debug_onoff = False
profile_file_path = "E:/工作/任务剖面/论文编写/剖面构造_V0.91-0712_指挥文电发送.xls"
save_dir = "./"
graph_name = "剖面构造_V0.91-0712_指挥文电发送"
state_func_color_dict = {}
path_state_nodes_list = []
path_state_probs_list = [] # 路径概率
test_cases_list = []
covered_path_states_set=set()
covered_path_trans_set=set()
covered_path_vars_set=set()
# ###############################################################
def read_profile_data(file_path:str)->list:
    ret = True
    profile_datas = []
    if os.path.exists(file_path) is False:
        print("剖面文件(%s)不存在!" % file_path)
        ret = False
    #
    if ret and file_path.endswith("xls") is False:
        print("剖面文件(%s)不是Excel(xls格式)文件!" % file_path)
        ret = False
    #
    if ret is True:
        try:
            profile_book = xlrd.open_workbook(file_path) #type:Book
            # 将xls中Sheet表名字保存到profile_datas
            names_data = {"sheet_names":list(profile_book.sheet_names())}
            profile_datas.append(pd.DataFrame(data=names_data))
        except Exception as e:
            print("打开%s失败(%s)!" % (file_path, str(e)))
            ret = False
        #
    #
    if ret is True:
        for name in profile_book.sheet_names():
            data = pd.read_excel(file_path, sheet_name=name, header=0)
            profile_datas.append(data)
            if debug_onoff is True:
                print("%s:" % name)
                print(data.values)
            #
        #
    #
    return profile_datas    
#

def build_profile_graph(profile_datas:list, graph_save_dir:str, graph_name:str) -> DiGraph:
    ret = True
    state_count = get_state_count(profile_datas)
    state_trans_prob_matrix = np.zeros((state_count, state_count),dtype=np.float32)
    state_trans_nx = DiGraph()
    state_trans_nx.add_nodes_from(range(1,state_count+1), VE=0, VL=0)
    profile_graph = Digraph(name = graph_name,encoding = 'utf-8')
    profile_graph.graph_attr.update(remincross="True") #最小交叉属性？
    if len(profile_datas)<1:
        ret = False
    #
    state_trans_sheet_index = get_sheet_index(profile_datas, "状态转移关系")
    if state_trans_sheet_index <=0:
        ret = False
    #
    
    if ret is True:
        if state_trans_sheet_index  < len(profile_datas) and state_trans_sheet_index>0:
            state_trans_data = profile_datas[state_trans_sheet_index] #type:pd.DataFrame
            for i in range(len(state_trans_data)):
                if len(state_trans_data.columns) >= 4:
                    this_state_func = get_state_func(profile_datas, state_trans_data.iloc[i,1])
                    this_state_color = get_state_func_color(this_state_func)
                    this_state_name = get_state_name(profile_datas, state_trans_data.iloc[i,1])                    
                    next_state_func = get_state_func(profile_datas, state_trans_data.iloc[i,2])
                    next_state_color = get_state_func_color(next_state_func)
                    next_state_name = get_state_name(profile_datas, state_trans_data.iloc[i,2])
                    trans_oe_name = get_oe_name(profile_datas, state_trans_data.iloc[i,3])
                    trans_oe_prob = get_oe_prob(profile_datas, state_trans_data.iloc[i,3]) 
                    # 创建graphviz图                 
                    profile_graph.node(name = this_state_name,color = this_state_color, fontname="FangSong")
                    profile_graph.node(name = next_state_name,color = next_state_color, fontname="FangSong")
                    profile_graph.edge(this_state_name,next_state_name, \
                        color = "green", label=trans_oe_name + "\n" + trans_oe_prob, \
                        fontname="FangSong") 
                    # 设置状态转移概率矩阵 
                    state_trans_prob_matrix[\
                        get_state_seq_no(profile_datas, state_trans_data.iloc[i,1]), \
                        get_state_seq_no(profile_datas, state_trans_data.iloc[i,2])] = float(trans_oe_prob)
                    # 设置状态转移网络图
                    state_trans_nx.add_weighted_edges_from([(\
                        get_state_seq_no(profile_datas, state_trans_data.iloc[i,1])+1, \
                        get_state_seq_no(profile_datas, state_trans_data.iloc[i,2])+1, \
                        float(trans_oe_prob)
                        )])
                #
            #
        #
    #
    if ret is True:
        # 保存graphviz图
        try:
            profile_graph.view(filename = graph_name, directory = save_dir, \
                quiet = True, quiet_view = True) # 自动显示并保存
            print("创建剖面图正常, 文件[%s(pdf及gv格式)]已保存到%s" % (graph_name, os.path.abspath(save_dir)))
        except Exception as e:
            print("创建剖面图发生异常(%s)!" % str(e))
            ret = False
        #
        # 保存状态转移概率矩阵 
        try:
            state_trans_prob_matrix_data = pd.DataFrame(state_trans_prob_matrix, \
                index = range(state_trans_prob_matrix.shape[0]),\
                columns = range(state_trans_prob_matrix.shape[1]))
            state_trans_prob_matrix_data.to_csv(os.path.join(save_dir,graph_name+".csv"), index=False)
            print("保存状态转移概率矩阵正常, 文件[%s]已保存到%s" % (graph_name+".csv", os.path.abspath(save_dir)))
        except Exception as e:
            print("保存状态转移概率矩阵异常(%s)!" % str(e))
            ret = False
        #        
    #
    return state_trans_nx
#

def search_test_paths(graph:DiGraph)->list:
    test_paths = []
    ret = True  
    state_nodes = list(graph.nodes)
    state_nodes.sort(reverse=False)
    if len(state_nodes) < 1:
        ret = False
    #
    if ret is True:
        start_node = state_nodes[0]
        end_node = state_nodes[-1]
        # 递归获取起始结点到终止结点的路径结点
        print("开始搜索状态转移测试路径......")
        search_path_state_nodes(graph, start_node, list(graph.neighbors(start_node)), end_node, [start_node])
        print("开始状态结点%d到结束状态结点%d共有%d条无环测试路径(结点序号从1开始)!" % (start_node, end_node, len(path_state_nodes_list)))    #
    #
    return test_paths
#

def search_path_state_nodes(graph:DiGraph, this_node:int, this_node_neighbors:list, end_node:int, path_nodes:list):
    if len(path_nodes)>(2*graph.number_of_nodes()):
        # 如果路径长度大于节点个数的2倍数
        return
    #

    if end_node not in this_node_neighbors:
        no_loop_node_neighbors = list(set(this_node_neighbors)-set([this_node])-set(path_nodes))
        #loop_node_neighbors = list(set(this_node_neighbors))
        #no_loop_node_neighbors = loop_node_neighbors
        for node in no_loop_node_neighbors:
            if node not in path_nodes or path_nodes.count(node)<10:#path_nodes.count(node)<3 合并环处理
                nodes = list()
                nodes.extend(path_nodes)
                nodes.append(node)
                search_path_state_nodes(graph, node, list(graph.neighbors(node)), end_node, nodes)
            else:
                # 有环，考虑环的处理策略
                pass
            #
        #
    else:
        nodes = list()        
        nodes.extend(path_nodes)
        nodes.append(end_node)
        global path_state_nodes_list
        if nodes not in path_state_nodes_list:           
            path_state_nodes_list.append(nodes)
            print(nodes,end=" ")
            # 获取路径概率  G[0][1]["weight"]
            probs = list()
            probs_sum = 0
            for i in range(len(nodes)-1):
                edge_prob = graph[nodes[i]][nodes[i+1]]["weight"]
                probs_sum = probs_sum + edge_prob
                probs.append(edge_prob)
            # end for
            path_state_probs_list.append(probs)
            print(",路径概率:" ,end="")
            print(probs,end=" ")
            print(",路径概率和:%f" % probs_sum)                        
        #
        no_loop_node_neighbors = list(set(this_node_neighbors)-set([this_node])-set(path_nodes))
        #loop_node_neighbors = list(set(this_node_neighbors))
        #no_loop_node_neighbors = loop_node_neighbors
        no_end_node_neighbors = list(set(no_loop_node_neighbors)-set([end_node]))
        for node in no_end_node_neighbors:
            '''
            nodes = list()        
            nodes.extend(path_nodes)
            nodes.append(node)
            search_path_state_nodes(graph, node, list(graph.neighbors(node)), end_node, nodes)
            '''
            if node not in path_nodes or path_nodes.count(node)<10:#path_nodes.count(node)<3 合并环处理
                nodes = list()
                nodes.extend(path_nodes)
                nodes.append(node)
                search_path_state_nodes(graph, node, list(graph.neighbors(node)), end_node, nodes)
            else:
                # 有环，考虑环的处理策略
                pass
            #
        #
    #
#

def extract_test_case_spec(profile_datas:list)-> bool:
    ret = True
    if len(path_state_nodes_list) < 1:
        ret = False
    #
    user_sheet_index = get_sheet_index(profile_datas, "用户")
    usesys_sheet_index = get_sheet_index(profile_datas, "用户_系统模式")
    func_sheet_index = get_sheet_index(profile_datas, "功能")
    sysfunc_sheet_index = get_sheet_index(profile_datas, "用户系统模式_功能")
    state_sheet_index = get_sheet_index(profile_datas, "系统特征状态")
    trans_sheet_index = get_sheet_index(profile_datas, "状态转移关系")
    oedef_sheet_index = get_sheet_index(profile_datas, "操作事件定义")
    oevar_sheet_index = get_sheet_index(profile_datas, "操作事件_变量")
    vardef_sheet_index = get_sheet_index(profile_datas, "变量定义")
    if user_sheet_index <= 0 or \
        usesys_sheet_index <= 0 or \
        func_sheet_index <= 0 or \
        sysfunc_sheet_index <= 0 or \
        state_sheet_index <= 0 or \
        trans_sheet_index <= 0 or \
        oedef_sheet_index <= 0 or \
        oevar_sheet_index <= 0 or \
        vardef_sheet_index <= 0 :
        ret = False
    #
    if ret is True:
        for path_state_nodes in path_state_nodes_list:
            if len(path_state_nodes) < 2:
                continue
            #
            test_case = TestCase()
            for i in range(len(path_state_nodes)-1):
                # 注意图中结点从1开始，DataFrame中下标是从0开始
                start_seq_no = path_state_nodes[i]-1
                end_seq_no = path_state_nodes[i+1]-1
                start_state_id = DataFrame(profile_datas[state_sheet_index]).iloc[start_seq_no,0]
                end_state_id = DataFrame(profile_datas[state_sheet_index]).iloc[end_seq_no,0]
                start_state_name = DataFrame(profile_datas[state_sheet_index]).iloc[start_seq_no,1]
                end_state_name = DataFrame(profile_datas[state_sheet_index]).iloc[end_seq_no,1]
                start_func_id = DataFrame(profile_datas[state_sheet_index]).iloc[start_seq_no,2]
                end_func_id = DataFrame(profile_datas[state_sheet_index]).iloc[end_seq_no,2]
                func_df=DataFrame(profile_datas[func_sheet_index])
                start_func_name = DataFrame(func_df.loc[func_df["功能编号"]==start_func_id,["功能名称"]]).values[0][0]
                end_func_name = DataFrame(func_df.loc[func_df["功能编号"]==end_func_id,["功能名称"]]).values[0][0]
                trans_df=DataFrame(profile_datas[trans_sheet_index])
                oe_id = DataFrame(
                    trans_df.loc[((trans_df["当前状态"]==start_state_id)&(trans_df["紧后状态"]==end_state_id)),\
                    ["操作/事件编号"]]).values[0][0]
                trans_id = DataFrame(
                    trans_df.loc[((trans_df["当前状态"]==start_state_id)&(trans_df["紧后状态"]==end_state_id)),\
                    ["转移关系编号"]]).values[0][0]
                oedef_df=DataFrame(profile_datas[oedef_sheet_index])
                oe_name = DataFrame(oedef_df.loc[oedef_df["操作/事件编号"]==oe_id, ["操作/事件名称"]]).values[0][0]
                oevar_df = DataFrame(profile_datas[oevar_sheet_index])
                oevar_data = DataFrame(oevar_df.loc[oevar_df["操作/事件编号"]==oe_id])                
                vardef_df = DataFrame(profile_datas[vardef_sheet_index])
                step_expect_list = list()
                for j in range(len(oevar_data)):
                    var_id = oevar_data.iloc[j,2]
                    var_attr = oevar_data.iloc[j,3]
                    var_value = oevar_data.iloc[j,4]
                    var_name = DataFrame(vardef_df.loc[vardef_df["变量编号"]==var_id, ["变量名称"]]).values[0][0]
                    step_expect_list.append("%s的%s变为%s" % (var_name, var_attr, str(var_value)))
                    # 记录覆盖情况(变量)
                    covered_path_vars_set.add(var_id) 
                #
                test_step = TestStep(oe_name, ";".join(step_expect_list) + "。")
                test_case.addState((start_state_name, start_func_name))
                test_case.addStep(test_step)
                if i == len(path_state_nodes)-2:
                    #添加最后1个state
                    test_case.addState((end_state_name, end_func_name))
                #
                # 记录覆盖情况(状态及状态对)
                covered_path_states_set.add(start_state_id)
                covered_path_states_set.add(end_state_id)
                covered_path_trans_set.add(trans_id)                                   
            # end for
            #保存测试用例
            test_cases_list.append(test_case)
        # end for
    #
    print("完成%d个测试序列的描述提取." % len(test_cases_list))
    return ret
#

def save_test_cases():
    ret = True
    try:        
        word_doc = docx.Document() #type:Document
        word_doc.add_paragraph()
    except Exception as e:
        print("新建word文件打开失败(%s)!" % (str(e)))
        ret = False
        return ret
    #
    for i in range(len(test_cases_list)):
        test_case = test_cases_list[i] #type:TestCase
        #创建Document Table   
        word_doc.add_heading("%d %s%d" % (i+1, "测试用例序列", i+1), 1) #type:Paragraph
        document_table = word_doc.add_table(len(test_case.steps) + 1,1) #type:Table
        document_table.style = "Table Grid"
        document_table.alignment = 0
        word_doc.add_paragraph(" ")
        test_states = ["%d(%s,%s)"%(test_case.states.index(state)+1, state[0],state[1]) for state in test_case.states]
        document_table.cell(0,0).text = "测试状态序列:\n" + "\n".join(test_states)
        document_table.cell(0,0).vertical_alignment = 1        
        for j in range(len(test_case.steps)):
            test_step = test_case.steps[j] # type:TestStep
            document_table.cell(j+1,0).text = ("步骤%d:" % (j+1)) + "\n" \
                + "操作:" + test_step.input + "。\n" \
                + "期望:" + test_step.expect
            #                   
            document_table.cell(j+1,0).vertical_alignment = 1            
        #
    #
    word_file_path = os.path.join(save_dir,  "测试用例序列("+graph_name+").docx")
    if os.path.exists(word_file_path):
        os.remove(word_file_path)
    #
    word_doc.save(word_file_path)
    print("测试用例序列已保存到%s." % word_file_path)
    #
#

def extract_sufficiency_indicator(profile_datas:list, graph:DiGraph)->OrderedDict:
    # 全部/关键特征状态总数/覆盖数
    # 全部/关键状态转移对总数
    # 全部/关键特征状态空间总数
    vardef_sheet_index = get_sheet_index(profile_datas, "变量定义")
    vardef_df = DataFrame(profile_datas[vardef_sheet_index])
    #
    indicator_dict = OrderedDict()
    indicator_dict["特征状态总数"]=len(graph.nodes)
    indicator_dict["状态转移对总数"]=len(graph.edges)
    indicator_dict["特征状态空间总数"]=len(vardef_df)
    indicator_dict["特征状态覆盖数"]=len(covered_path_states_set)
    indicator_dict["状态转移对覆盖数"]=len(covered_path_trans_set)
    indicator_dict["特征状态空间覆盖数"]=len(covered_path_vars_set)
    indicator_dict["关键特征状态总数"]=len(graph.nodes)
    indicator_dict["关键状态转移对总数"]=len(graph.edges)
    indicator_dict["关键特征状态空间总数"]=len(vardef_df)
    indicator_dict["关键特征状态覆盖数"]=len(covered_path_states_set)
    indicator_dict["关键状态转移对覆盖数"]=len(covered_path_trans_set)
    indicator_dict["关键特征状态空间覆盖数"]=len(covered_path_vars_set)
    #保存
    indicator_df =DataFrame.from_dict(indicator_dict, orient='index', columns=["度量值"])
    csv_file_path = os.path.join(save_dir,  "度量项及取值("+graph_name+").csv")
    indicator_df.to_csv(csv_file_path, index=True, encoding="utf_8_sig")
    print("度量项及取值已保存到%s." % csv_file_path)
    #
    return indicator_dict
#

def init_state_func_color(profile_datas:list) -> bool:
    sheet_name = "功能"
    func_id_list = []
    ret = True
    func_sheet_index = get_sheet_index(profile_datas, sheet_name)
    if func_sheet_index <= 0:
        ret = False
    #
    if ret is True:
        if func_sheet_index  < len(profile_datas) and func_sheet_index>0:
            state_define_data = profile_datas[func_sheet_index] #type:pd.DataFrame
            try:
                func_id_list = list(state_define_data["功能编号"])
            #
            except Exception as e:
                ret = False
                print("在%s中查找到%s发生异常(%s)!" % (sheet_name, "功能编号", str(e)))
            #
        #
    #
    my_color_keys = list(my_color_dict.keys())
    for i in range(len(func_id_list)):
        state_func_color_dict[func_id_list[i]] = my_color_keys[i% len(my_color_keys)]        
    #
    return ret
#


def get_state_func_color(func_id:str)->str:
    func_color = "red"
    try:
        func_color_key = state_func_color_dict[func_id]
        func_color = my_color_dict[func_color_key]
    except Exception as e:
        print("查找%s功能对应的颜色发生异常(%s)!" % (func_id, str(e)))
    #
    return func_color
#

def get_state_count(profile_datas:list)->str:
    sheet_name = "系统特征状态"
    state_count = 0
    ret = True
    state_define_sheet_index = get_sheet_index(profile_datas, sheet_name)
    if state_define_sheet_index <= 0:
        ret = False
    #
    if ret is True:
        if state_define_sheet_index  < len(profile_datas) and state_define_sheet_index>0:
            state_define_data = profile_datas[state_define_sheet_index] #type:pd.DataFrame
            state_count = len(state_define_data)
        #
    #
    return state_count
#

def get_vars_count(profile_datas:list)->str:
    sheet_name = "变量定义"
    var_count = 0
    ret = True
    var_define_sheet_index = get_sheet_index(profile_datas, sheet_name)
    if var_define_sheet_index <= 0:
        ret = False
    #
    if ret is True:
        if var_define_sheet_index  < len(profile_datas) and var_define_sheet_index>0:
            var_define_data = profile_datas[var_define_sheet_index] #type:pd.DataFrame
            var_count = len(var_define_data)
        #
    #
    return var_count
#

def get_state_func(profile_datas:list, state_id:str)->str:
    sheet_name = "系统特征状态"
    fund_id = ""
    ret = True
    state_define_sheet_index = get_sheet_index(profile_datas, sheet_name)
    if state_define_sheet_index <= 0:
        ret = False
    #
    if ret is True:
        if state_define_sheet_index  < len(profile_datas) and state_define_sheet_index>0:
            state_define_data = profile_datas[state_define_sheet_index] #type:pd.DataFrame
            try:
                this_state_func_data = \
                    state_define_data.loc[state_define_data["系统特征状态编号"]==state_id, ["所在功能"]]
                if len(this_state_func_data)>0:
                    fund_id = this_state_func_data.values[0][0]
                #
            #
            except Exception as e:
                print("在%s中查找到%s发生异常(%s)!" % (sheet_name, state_id, str(e)))
            #
        #
    #
    return fund_id
#

def get_state_name(profile_datas:list, state_id:str)->str:
    sheet_name = "系统特征状态"
    state_name = state_id
    ret = True
    state_define_sheet_index = get_sheet_index(profile_datas, sheet_name)
    if state_define_sheet_index <= 0:
        ret = False
    #
    if ret is True:
        if state_define_sheet_index  < len(profile_datas) and state_define_sheet_index>0:
            state_define_data = profile_datas[state_define_sheet_index] #type:pd.DataFrame
            try:
                this_state_name_data = \
                    state_define_data.loc[state_define_data["系统特征状态编号"]==state_id, ["系统特征状态名称"]]
                if len(this_state_name_data)>0:
                    state_name = this_state_name_data.values[0][0]
                #
            #
            except Exception as e:
                print("在%s中查找到%s发生异常(%s)!" % (sheet_name, state_id, str(e)))
            #
        #
    #
    return state_name
#

def get_state_seq_no(profile_datas:list, state_id:str)->int:
    sheet_name = "系统特征状态"
    state_seq_no = -1
    ret = True
    state_define_sheet_index = get_sheet_index(profile_datas, sheet_name)
    if state_define_sheet_index <= 0:
        ret = False
    #
    if ret is True:
        if state_define_sheet_index  < len(profile_datas) and state_define_sheet_index>0:
            state_define_data = profile_datas[state_define_sheet_index] #type:pd.DataFrame
            try:
                this_state_seq_index = \
                    state_define_data.loc[state_define_data["系统特征状态编号"]==state_id].index
                if len(list(this_state_seq_index))>0:
                    state_seq_no = list(this_state_seq_index)[0]
                else:                
                    ret = False
                #
            #
            except Exception as e:
                print("在%s中查找到%s发生异常(%s)!" % (sheet_name, state_id, str(e)))
            #     
        #
    #
    return state_seq_no
#

def get_oe_name(profile_datas:list, oe_id:str)->str:
    sheet_name = "操作事件定义"
    oe_name = oe_id
    ret = True
    oe_define_sheet_index = get_sheet_index(profile_datas, sheet_name)
    if oe_define_sheet_index <= 0:
        ret = False
    #
    if ret is True:
        if oe_define_sheet_index  < len(profile_datas) and oe_define_sheet_index>0:
            oe_define_data = profile_datas[oe_define_sheet_index] #type:pd.DataFrame
            try:
                this_oe_name_data = \
                    oe_define_data.loc[oe_define_data["操作/事件编号"]==oe_id, ["操作/事件名称"]]
                if len(this_oe_name_data)>0:
                    oe_name = this_oe_name_data.values[0][0]
                #
            #
            except Exception as e:
                print("在%s中查找到%s发生异常(%s)!" % (sheet_name, oe_id, str(e)))
            #
        #
    #
    return oe_name
#

def get_oe_prob(profile_datas:list, oe_id:str)->str:
    sheet_name = "操作事件定义"
    oe_prob = "0.00"
    ret = True
    oe_define_sheet_index = get_sheet_index(profile_datas, sheet_name)
    if oe_define_sheet_index <= 0:
        ret = False
    #
    if ret is True:
        if oe_define_sheet_index  < len(profile_datas) and oe_define_sheet_index>0:
            oe_define_data = profile_datas[oe_define_sheet_index] #type:pd.DataFrame
            try:
                this_oe_prob_data = \
                    oe_define_data.loc[oe_define_data["操作/事件编号"]==oe_id, ["操作/事件概率"]]
                if len(this_oe_prob_data)>0:
                    oe_prob = this_oe_prob_data.values[0][0]
                #
            #
            except Exception as e:
                print("在%s中查找到%s发生异常(%s)!" % (sheet_name, oe_id, str(e)))
            #
        #
    #
    return str(oe_prob)
#

def get_sheet_index(profile_datas:list, sheet_name:str) -> int:
    ret = True
    sheet_index = -1
    if len(profile_datas)<1:
        ret = False
    #    
    if ret is True:
        sheet_names = profile_datas[0] #type:pd.DataFrame        
        try:
            sheet_index = sheet_names.loc[sheet_names["sheet_names"]==sheet_name].index
            if len(list(sheet_index))>0:
                sheet_index = list(sheet_index)[0]
            else:                
                ret = False
            #
        #
        except Exception as e:
            print("在剖面数据中未查找到%s(%s)!" % (sheet_name, str(e)))
        #
    #
    # 注意profile_datas第0个表是表名称列表
    # 所以在profile_datas中状态转移关系的表位置还要+1
    if sheet_index>=0:
        sheet_index = sheet_index + 1
    else:
        sheet_index = -1
        print("在剖面数据中未查找到%s!" % sheet_name)
    #
    return sheet_index
#

class TestStep():
    def __init__(self,input:str,expect:str):
        self.__input = input
        self.__expect = expect
    #

    @property
    def input(self):
        return self.__input
    #

    @input.setter
    def input(self,input:str):
        self.__input = input
    #

    @property
    def expect(self):
        return self.__expect
    #

    @expect.setter
    def expect(self,expect:str):
        self.__expect = expect
    #
#

#测试用例
class TestCase():
    def __init__(self):
        self.__states = []
        self.__testSteps = []
    #

    def addState(self, state:tuple):
        #state:(state_name, func_name)
        self.__states.append(state)
    #
    @property
    def states(self):
        return self.__states
    #

    def addStep(self, step:TestStep):
        self.__testSteps.append(step)
    #
    @property
    def steps(self):
        return self.__testSteps
    #
#

def main():
    print("Hello, my mission profile...")
    # 读取剖面excel数据
    profile_datas = read_profile_data(profile_file_path)
    # 初始化颜色表
    init_state_func_color(profile_datas)
    # 构建剖面网络图（附带graphviz图文件，状态转移概率矩阵文件）
    state_trans_nx = build_profile_graph(profile_datas, save_dir, graph_name)
    # 输出初始结点到结束结点之间的所有不重复可达路径，注意环的处理        
    search_test_paths(state_trans_nx)
    # 提取测试用例描述
    extract_test_case_spec(profile_datas)
    # 保存测试用例
    save_test_cases()
    # 获取度量项及取值并保存
    extract_sufficiency_indicator(profile_datas, state_trans_nx)
    #
    print("已覆盖的特征状态数:%d(总数%d)" % (len(covered_path_states_set), state_trans_nx.number_of_nodes()))
    print(covered_path_states_set)
    print("已覆盖的转移边数:%d(总数%d)" % (len(covered_path_trans_set), state_trans_nx.number_of_edges()))
    print(covered_path_trans_set)
    print("已覆盖的变量数:%d(总数%d)" % (len(covered_path_vars_set), get_vars_count(profile_datas)))
    print(covered_path_vars_set)
    return 0
#

if __name__ == "__main__":
    os.system("chcp 65001")
    ret = main() #type:int
    sys.exit(ret)
#